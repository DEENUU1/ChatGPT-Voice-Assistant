import asyncio
from docx import Document
from summaryzer import return_article_summary
import configparser
from welcome import Welcome
from text_to_speech import return_speech
from speech_command import speech_command
from openai_async import get_openai_conclusion
from weather import Weather
from image_converter import image_to_text


async def main_async():
    # Configuration
    config = configparser.ConfigParser()
    config.read('config.ini')
    country_code = config.get('NEWS SUMMARIZER', 'country_code')
    user_name = config.get('NEWS SUMMARIZER', 'user_name')
    document = Document()
    welcome = Welcome()
    welcome_user = welcome.return_welcome_message(user_name)

    # Welcome massage
    print(welcome_user)
    return_speech(welcome_user)
    print("I am listening... ")
    # Options based on the user decision
    while True:
        print("1. Weather information for today SAY 'WEATHER'"
              "\n2. News summaries SAY 'NEWS'"
              "\n3. AI Conversation SAY 'CONVERSATION'"
              "\n4. Explain code SAY 'EXPLAIN'"
              "\n5. Image -> Text SAY 'IMAGE'")
        user_decision = speech_command()

        # Weather info
        if "weather" in user_decision:
            weather = Weather()
            weather_summary = weather.return_weather_summary()
            print(await get_openai_conclusion(weather_summary))
            return_speech(await get_openai_conclusion(weather_summary))
            continue

        # News summaries
        if "news" in user_decision:
            async for summary, article_url in return_article_summary(country_code):
                print(summary, article_url)
                return_speech(summary)
                document.add_paragraph(
                    summary + ' ' + article_url)
            document.save('news.docx')
            continue

        # Conversation with AI
        if "conversation" in user_decision:
            while True:
                print("I am listening... ")
                prompt = speech_command()
                if "close the program" in prompt:
                    break
                print(await get_openai_conclusion(prompt))
                return_speech(await get_openai_conclusion(prompt))

        if "explain" in user_decision:
            while True:
                # Adding multiple lines of code
                content = []
                while True:
                    try:
                        command = str(input("> "))
                        if command == "###":
                            break
                    except EOFError:
                        break
                    content.append(command)

                prompt = f"explain me this code {content}"
                print(await get_openai_conclusion(prompt))
                return_speech(await get_openai_conclusion(prompt))
                break

        # Image ot text converter and AI notes
        if "image" in user_decision:
            while True:
                image_name = input("Name of the image > ")
                what_to_do = input("Tell me what to do with this text > ")
                language = input("In what language is this text? > ")
                data = image_to_text(image_name, language)
                prompt = f"Summarize me this text, {what_to_do} {data}"
                openai_summary = await get_openai_conclusion(prompt)
                print(openai_summary)
                return_speech(openai_summary)
                document.add_paragraph(openai_summary)
                document.save('imageFromText.docx')
                break

if __name__ == '__main__':
    asyncio.run(main_async())
